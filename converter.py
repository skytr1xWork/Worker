import csv
import html
import io
import json
import os
import re
import subprocess
import tempfile
from docx import Document
import markdown
from PIL import Image

# ==========================================
# Image Formats
# ==========================================
SUPPORTED_IMAGE_FORMATS = {
    "PNG": {"ext": "png", "mime": "image/png"},
    "JPG": {"ext": "jpg", "mime": "image/jpeg"},
    "WEBP": {"ext": "webp", "mime": "image/webp"},
    "BMP": {"ext": "bmp", "mime": "image/bmp"},
    "TIFF": {"ext": "tiff", "mime": "image/tiff"},
    "ICO": {"ext": "ico", "mime": "image/x-icon"},
    "PDF": {"ext": "pdf", "mime": "application/pdf"},
    "GIF": {"ext": "gif", "mime": "image/gif"},
}

IMAGE_EXTENSIONS = {
    "png": "PNG",
    "jpg": "JPG",
    "jpeg": "JPG",
    "webp": "WEBP",
    "bmp": "BMP",
    "tiff": "TIFF",
    "tif": "TIFF",
    "ico": "ICO",
    "gif": "GIF",
    "pdf": "PDF",
}

IMAGE_MIME_TYPES = {
    "image/png": "PNG",
    "image/jpeg": "JPG",
    "image/pjpeg": "JPG",
    "image/webp": "WEBP",
    "image/bmp": "BMP",
    "image/x-ms-bmp": "BMP",
    "image/tiff": "TIFF",
    "image/x-icon": "ICO",
    "image/vnd.microsoft.icon": "ICO",
    "image/gif": "GIF",
}

# ==========================================
# Document Formats
# ==========================================
SUPPORTED_DOCUMENT_FORMATS = {
    "DOCX": {"ext": "docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "MD": {"ext": "md", "mime": "text/markdown"},
    "TXT": {"ext": "txt", "mime": "text/plain"},
    "DAT": {"ext": "dat", "mime": "application/octet-stream"},
    "CSV": {"ext": "csv", "mime": "text/csv"},
    "TSV": {"ext": "tsv", "mime": "text/tab-separated-values"},
    "JSON": {"ext": "json", "mime": "application/json"},
    "XML": {"ext": "xml", "mime": "application/xml"},
    "LOG": {"ext": "log", "mime": "text/plain"},
    "HTML": {"ext": "html", "mime": "text/html"},
}

DOCUMENT_EXTENSIONS = {
    "docx": "DOCX",
    "md": "MD",
    "markdown": "MD",
    "txt": "TXT",
    "dat": "DAT",
    "csv": "CSV",
    "tsv": "TSV",
    "json": "JSON",
    "xml": "XML",
    "log": "LOG",
    "html": "HTML",
    "htm": "HTML",
}

DOCUMENT_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
    "text/markdown": "MD",
    "text/x-markdown": "MD",
    "text/plain": "TXT",
    "text/csv": "CSV",
    "text/tab-separated-values": "TSV",
    "application/json": "JSON",
    "application/xml": "XML",
    "text/xml": "XML",
    "text/html": "HTML",
}

DOCUMENT_TARGETS = {
    "DOCX": ["MD", "TXT", "HTML", "JSON", "DAT", "LOG"],
    "MD": ["DOCX", "HTML", "JSON", "TXT", "DAT", "LOG"],
    "TXT": ["JSON", "HTML", "DOCX", "MD", "DAT", "LOG", "CSV"],
    "DAT": ["JSON", "HTML", "TXT", "DOCX", "MD", "LOG", "CSV"],
    "LOG": ["JSON", "HTML", "TXT", "DOCX", "MD", "DAT"],
    "CSV": ["JSON", "HTML", "MD", "DOCX", "TSV", "TXT", "DAT"],
    "TSV": ["JSON", "HTML", "CSV", "MD", "DOCX", "TXT", "DAT"],
    "JSON": ["CSV", "HTML", "TXT", "MD", "DOCX", "DAT"],
    "XML": ["JSON", "HTML", "TXT", "MD", "DOCX", "DAT"],
    "HTML": ["MD", "TXT", "DOCX", "JSON", "DAT", "LOG"],
}

# ==========================================
# Audio Formats
# ==========================================
SUPPORTED_AUDIO_FORMATS = {
    "MP3": {"ext": "mp3", "mime": "audio/mpeg"},
    "WAV": {"ext": "wav", "mime": "audio/wav"},
    "OGG": {"ext": "ogg", "mime": "audio/ogg"},
    "OPUS": {"ext": "opus", "mime": "audio/opus"},
    "FLAC": {"ext": "flac", "mime": "audio/flac"},
    "AAC": {"ext": "aac", "mime": "audio/aac"},
    "M4A": {"ext": "m4a", "mime": "audio/mp4"},
    "WMA": {"ext": "wma", "mime": "audio/x-ms-wma"},
    "AIFF": {"ext": "aiff", "mime": "audio/aiff"},
    "AMR": {"ext": "amr", "mime": "audio/amr"},
    "AC3": {"ext": "ac3", "mime": "audio/ac3"},
    "MP2": {"ext": "mp2", "mime": "audio/mp2"},
}

AUDIO_EXTENSIONS = {
    "mp3": "MP3",
    "wav": "WAV",
    "ogg": "OGG",
    "opus": "OPUS",
    "oga": "OPUS",
    "flac": "FLAC",
    "aac": "AAC",
    "m4a": "M4A",
    "wma": "WMA",
    "aiff": "AIFF",
    "aif": "AIFF",
    "amr": "AMR",
    "ac3": "AC3",
    "mp2": "MP2",
}

AUDIO_MIME_TYPES = {
    "audio/mpeg": "MP3",
    "audio/mp3": "MP3",
    "audio/wav": "WAV",
    "audio/x-wav": "WAV",
    "audio/ogg": "OGG",
    "audio/opus": "OPUS",
    "audio/x-opus+ogg": "OPUS",
    "audio/flac": "FLAC",
    "audio/x-flac": "FLAC",
    "audio/aac": "AAC",
    "audio/mp4": "M4A",
    "audio/x-m4a": "M4A",
    "audio/x-ms-wma": "WMA",
    "audio/aiff": "AIFF",
    "audio/x-aiff": "AIFF",
    "audio/amr": "AMR",
    "audio/ac3": "AC3",
}


def normalize_format(fmt: str) -> str:
    """Normalizes format string to standard uppercase key."""
    if not fmt:
        return ""
    cleaned = fmt.upper().strip().lstrip(".")
    if cleaned in ("JPEG", "JPG"):
        return "JPG"
    if cleaned in ("TIF", "TIFF"):
        return "TIFF"
    if cleaned == "MARKDOWN":
        return "MD"
    if cleaned == "HTM":
        return "HTML"
    if cleaned in ("AIF", "AIFF"):
        return "AIFF"
    if cleaned == "OGA":
        return "OPUS"
    return cleaned


def detect_file_type(filename: str | None = None, mime_type: str | None = None) -> tuple[str | None, str | None]:
    """
    Detects file category ('image', 'document', 'audio') and format name.
    """
    ext = os.path.splitext(filename)[1].lower().lstrip(".") if filename else ""

    # 1. Check audio extension
    if ext in AUDIO_EXTENSIONS:
        return "audio", AUDIO_EXTENSIONS[ext]

    # 2. Check image extension
    if ext in IMAGE_EXTENSIONS:
        return "image", IMAGE_EXTENSIONS[ext]

    # 3. Check document extension
    if ext in DOCUMENT_EXTENSIONS:
        return "document", DOCUMENT_EXTENSIONS[ext]

    # Check MIME types
    if mime_type:
        mime = mime_type.lower()

        # Audio MIME
        if mime.startswith("audio/"):
            for m_key, f_val in AUDIO_MIME_TYPES.items():
                if m_key in mime:
                    return "audio", f_val
            if "opus" in mime or "oga" in mime:
                return "audio", "OPUS"
            if "ogg" in mime:
                return "audio", "OGG"
            if "flac" in mime:
                return "audio", "FLAC"
            if "wav" in mime:
                return "audio", "WAV"
            return "audio", "MP3"

        # Image MIME
        if mime.startswith("image/"):
            for m_key, f_val in IMAGE_MIME_TYPES.items():
                if m_key in mime:
                    return "image", f_val
            if "png" in mime:
                return "image", "PNG"
            if "jpeg" in mime or "jpg" in mime:
                return "image", "JPG"
            if "webp" in mime:
                return "image", "WEBP"
            return "image", "PNG"

        # Document MIME
        if mime in DOCUMENT_MIME_TYPES:
            return "document", DOCUMENT_MIME_TYPES[mime]
        if "wordprocessingml" in mime:
            return "document", "DOCX"
        if "json" in mime:
            return "document", "JSON"
        if "csv" in mime:
            return "document", "CSV"
        if "xml" in mime:
            return "document", "XML"
        if "html" in mime:
            return "document", "HTML"
        if "text" in mime:
            return "document", "TXT"

    return None, None


def detect_image_format(filename: str | None = None, mime_type: str | None = None) -> str | None:
    cat, fmt = detect_file_type(filename, mime_type)
    return fmt if cat == "image" else None


def format_size(size_bytes: int) -> str:
    """Formats file size into human-readable string."""
    if size_bytes < 1024:
        return f"{size_bytes} Б"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} КБ"
    return f"{size_bytes / (1024 * 1024):.2f} МБ"


def decode_text(raw_bytes: bytes) -> str:
    """Safely decodes raw bytes into string trying multiple encodings."""
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


# ==========================================
# Audio Converter (using FFmpeg with high compatibility)
# ==========================================

def convert_audio(input_bytes: bytes, source_format: str, target_format: str, orig_filename: str | None = None) -> tuple[bytes, str]:
    """
    Converts audio raw bytes into the requested target format using FFmpeg.
    Ensures audible volume, proper sample rates, and correct codecs.
    Returns (output_bytes, file_extension).
    """
    src = normalize_format(source_format)
    target = normalize_format(target_format)

    if target not in SUPPORTED_AUDIO_FORMATS:
        raise ValueError(f"Неподдерживаемый целевой аудиоформат: {target_format}")

    # Use original extension if available so FFmpeg knows the demuxer
    src_ext = "audio"
    if orig_filename and "." in orig_filename:
        src_ext = orig_filename.rsplit(".", 1)[-1].lower()
    elif src in SUPPORTED_AUDIO_FORMATS:
        src_ext = SUPPORTED_AUDIO_FORMATS[src]["ext"]

    dst_ext = SUPPORTED_AUDIO_FORMATS[target]["ext"]

    with tempfile.NamedTemporaryFile(suffix=f".{src_ext}", delete=False) as src_f:
        src_f.write(input_bytes)
        src_path = src_f.name

    with tempfile.NamedTemporaryFile(suffix=f".{dst_ext}", delete=False) as dst_f:
        dst_path = dst_f.name

    try:
        cmd = ["ffmpeg", "-y", "-i", src_path, "-vn"]
        if target == "MP3":
            cmd.extend(["-c:a", "libmp3lame", "-b:a", "192k", "-ar", "44100", "-ac", "2"])
        elif target == "OPUS":
            cmd.extend(["-c:a", "libopus", "-b:a", "128k", "-ar", "48000", "-vbr", "on"])
        elif target == "OGG":
            cmd.extend(["-c:a", "libvorbis", "-q:a", "4", "-ar", "44100"])
        elif target == "FLAC":
            cmd.extend(["-c:a", "flac", "-ar", "44100"])
        elif target in ("AAC", "M4A"):
            cmd.extend(["-c:a", "aac", "-b:a", "192k", "-ar", "44100", "-movflags", "+faststart"])
        elif target == "WAV":
            cmd.extend(["-c:a", "pcm_s16le", "-ar", "44100", "-ac", "2"])
        elif target == "AMR":
            cmd.extend(["-ar", "8000", "-ac", "1", "-c:a", "libopencore_amrnb"])
        elif target == "AIFF":
            cmd.extend(["-c:a", "pcm_s16be", "-ar", "44100"])
        elif target == "AC3":
            cmd.extend(["-c:a", "ac3", "-b:a", "192k", "-ar", "44100"])
        elif target == "MP2":
            cmd.extend(["-c:a", "mp2", "-b:a", "192k", "-ar", "44100"])
        elif target == "WMA":
            cmd.extend(["-c:a", "wmav2", "-b:a", "192k", "-ar", "44100"])

        cmd.append(dst_path)

        res = subprocess.run(cmd, capture_output=True, timeout=60)
        if res.returncode != 0:
            error_msg = res.stderr.decode("utf-8", errors="replace")
            raise RuntimeError(f"FFmpeg ошибка: {error_msg[-200:]}")

        with open(dst_path, "rb") as out_f:
            output_bytes = out_f.read()

        return output_bytes, dst_ext
    finally:
        if os.path.exists(src_path):
            os.unlink(src_path)
        if os.path.exists(dst_path):
            os.unlink(dst_path)


# ==========================================
# Image Converter
# ==========================================

def convert_image(input_bytes: bytes, target_format: str) -> tuple[bytes, str]:
    """
    Converts image raw bytes into the requested target format.
    Returns a tuple of (converted_bytes, file_extension).
    """
    target = normalize_format(target_format)
    if target not in SUPPORTED_IMAGE_FORMATS:
        raise ValueError(f"Неподдерживаемый целевой формат: {target_format}")

    with Image.open(io.BytesIO(input_bytes)) as img:
        output_io = io.BytesIO()
        ext = SUPPORTED_IMAGE_FORMATS[target]["ext"]

        if target == "JPG":
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                img_rgba = img.convert("RGBA")
                bg.paste(img_rgba, mask=img_rgba.split()[3])
                save_img = bg
            elif img.mode != "RGB":
                save_img = img.convert("RGB")
            else:
                save_img = img
            save_img.save(output_io, format="JPEG", quality=95, optimize=True)

        elif target == "PNG":
            img.save(output_io, format="PNG", optimize=True)

        elif target == "WEBP":
            img.save(output_io, format="WEBP", quality=95)

        elif target == "BMP":
            if img.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                img_rgba = img.convert("RGBA")
                bg.paste(img_rgba, mask=img_rgba.split()[3])
                save_img = bg
            elif img.mode != "RGB":
                save_img = img.convert("RGB")
            else:
                save_img = img
            save_img.save(output_io, format="BMP")

        elif target == "TIFF":
            img.save(output_io, format="TIFF")

        elif target == "ICO":
            ico_img = img.copy()
            if ico_img.width > 256 or ico_img.height > 256:
                ico_img.thumbnail((256, 256), Image.Resampling.LANCZOS)
            if ico_img.mode not in ("RGBA", "RGB"):
                ico_img = ico_img.convert("RGBA")
            ico_img.save(output_io, format="ICO")

        elif target == "PDF":
            if img.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                img_rgba = img.convert("RGBA")
                bg.paste(img_rgba, mask=img_rgba.split()[3])
                save_img = bg
            elif img.mode != "RGB":
                save_img = img.convert("RGB")
            else:
                save_img = img
            save_img.save(output_io, format="PDF", resolution=100.0)

        elif target == "GIF":
            img.save(output_io, format="GIF")

        return output_io.getvalue(), ext


# ==========================================
# Responsive HTML Page Template
# ==========================================

def _wrap_in_readable_html_template(body_html: str, title: str = "Документ") -> str:
    """Wraps HTML content in a modern, reader-friendly styling with signature footer."""
    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      --bg-color: #f8fafc;
      --card-bg: #ffffff;
      --text-primary: #0f172a;
      --text-secondary: #475569;
      --border-color: #e2e8f0;
      --accent-color: #2563eb;
      --code-bg: #f1f5f9;
      --table-header-bg: #f8fafc;
      --table-stripe-bg: #f8fafc;
    }}
    @media (prefers-color-scheme: dark) {{
      :root {{
        --bg-color: #0f172a;
        --card-bg: #1e293b;
        --text-primary: #f8fafc;
        --text-secondary: #94a3b8;
        --border-color: #334155;
        --accent-color: #38bdf8;
        --code-bg: #0f172a;
        --table-header-bg: #0f172a;
        --table-stripe-bg: #162032;
      }}
    }}
    * {{ box-sizing: border-box; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
      background-color: var(--bg-color);
      color: var(--text-primary);
      line-height: 1.75;
      margin: 0;
      padding: 40px 16px;
      -webkit-font-smoothing: antialiased;
    }}
    .document-card {{
      max-width: 820px;
      margin: 0 auto;
      background: var(--card-bg);
      padding: 48px;
      border-radius: 12px;
      box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.05), 0 2px 4px -2px rgb(0 0 0 / 0.05);
      border: 1px solid var(--border-color);
    }}
    @media (max-width: 640px) {{
      .document-card {{ padding: 24px 16px; }}
    }}
    h1, h2, h3, h4, h5, h6 {{
      color: var(--text-primary);
      font-weight: 600;
      margin-top: 1.5em;
      margin-bottom: 0.6em;
      line-height: 1.3;
    }}
    h1 {{ font-size: 2rem; border-bottom: 2px solid var(--border-color); padding-bottom: 0.3em; margin-top: 0; }}
    h2 {{ font-size: 1.5rem; border-bottom: 1px solid var(--border-color); padding-bottom: 0.25em; }}
    h3 {{ font-size: 1.25rem; }}
    p {{ margin-top: 0; margin-bottom: 1.2em; color: var(--text-primary); }}
    a {{ color: var(--accent-color); text-decoration: none; }}
    a:hover {{ text-decoration: underline; }}
    ul, ol {{ margin-top: 0; margin-bottom: 1.2em; padding-left: 24px; }}
    li {{ margin-bottom: 0.4em; }}
    pre, code {{
      font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace;
      font-size: 0.9em;
      background-color: var(--code-bg);
      border-radius: 6px;
    }}
    code {{ padding: 0.2em 0.4em; }}
    pre {{
      padding: 16px;
      overflow-x: auto;
      border: 1px solid var(--border-color);
      line-height: 1.5;
    }}
    pre code {{ padding: 0; background: transparent; }}
    blockquote {{
      border-left: 4px solid var(--accent-color);
      margin: 1.5em 0;
      padding: 0.6em 1.2em;
      color: var(--text-secondary);
      background-color: var(--table-stripe-bg);
      border-radius: 0 8px 8px 0;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 1.5em 0;
      font-size: 0.95em;
    }}
    th, td {{
      padding: 12px 14px;
      text-align: left;
      border: 1px solid var(--border-color);
    }}
    th {{
      background-color: var(--table-header-bg);
      font-weight: 600;
    }}
    tr:nth-child(even) {{
      background-color: var(--table-stripe-bg);
    }}
    .log-badge {{
      display: inline-block;
      padding: 2px 8px;
      font-size: 0.75rem;
      font-weight: 600;
      border-radius: 4px;
      text-transform: uppercase;
      margin-right: 6px;
    }}
    .log-info {{ background: #e0f2fe; color: #0369a1; }}
    .log-warn {{ background: #fef3c7; color: #b45309; }}
    .log-error {{ background: #fee2e2; color: #b91c1c; }}
    .log-debug {{ background: #f1f5f9; color: #475569; }}
    .log-row {{ font-family: monospace; font-size: 0.9em; padding: 6px 0; border-bottom: 1px solid var(--border-color); }}
    .log-ts {{ color: var(--text-secondary); margin-right: 8px; }}
    .document-footer {{
      margin-top: 40px;
      padding-top: 20px;
      border-top: 1px solid var(--border-color);
      font-size: 0.85rem;
      color: var(--text-secondary);
      text-align: center;
    }}
    .document-footer a {{
      color: var(--accent-color);
      font-weight: 500;
    }}
  </style>
</head>
<body>
  <div class="document-card">
    {body_html}
    <footer class="document-footer">
      сделано благодаря лучшему, красивому, офигенному, невероятному <a href="https://t.me/skytr1xHelper_bot" target="_blank" rel="noopener">@skytr1xHelper_bot</a>
    </footer>
  </div>
</body>
</html>"""


# ==========================================
# Structured JSON Converters
# ==========================================

def _infer_value_type(val: str):
    """Auto-converts string into int, float, bool, or original string."""
    v_clean = val.strip()
    if v_clean.lower() == "true":
        return True
    if v_clean.lower() == "false":
        return False
    if v_clean.lower() in ("null", "none"):
        return None
    if v_clean.isdigit() or (v_clean.startswith("-") and v_clean[1:].isdigit()):
        try:
            return int(v_clean)
        except ValueError:
            pass
    try:
        return float(v_clean)
    except ValueError:
        pass
    return val


def text_to_structured_json(text: str) -> str:
    """Intelligently parses plain text into structured, readable JSON."""
    try:
        parsed = json.loads(text)
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    except Exception:
        pass

    lines = text.splitlines()
    non_empty_lines = [l.strip() for l in lines if l.strip()]

    kv_pattern = re.compile(r'^([a-zA-Z0-9_\-\.\s]+)\s*[:=]\s*(.*)$')
    key_values = {}
    sections = {}
    current_section = "general"
    kv_count = 0

    for line in non_empty_lines:
        if line.startswith("[") and line.endswith("]"):
            current_section = line[1:-1].strip()
            if current_section not in sections:
                sections[current_section] = {}
            continue
        m = kv_pattern.match(line)
        if m:
            k, v = m.group(1).strip(), m.group(2).strip()
            val_cast = _infer_value_type(v)
            if current_section not in sections:
                sections[current_section] = {}
            sections[current_section][k] = val_cast
            key_values[k] = val_cast
            kv_count += 1

    log_pattern = re.compile(r'^(\d{4}[-/.]\d{2}[-/.]\d{2}[ T]\d{2}:\d{2}:\d{2}(?:[.,]\d+)?)\s*(?:\[(\w+)\]|(\w+):)?\s*(.*)$')
    log_entries = []
    for line in non_empty_lines:
        m = log_pattern.match(line)
        if m:
            ts, lvl1, lvl2, msg = m.groups()
            level = (lvl1 or lvl2 or "INFO").upper()
            log_entries.append({
                "timestamp": ts,
                "level": level,
                "message": msg.strip()
            })

    words = text.split()
    meta = {
        "lines_count": len(lines),
        "words_count": len(words),
        "characters_count": len(text)
    }

    if len(log_entries) >= len(non_empty_lines) * 0.7 and len(log_entries) > 0:
        result = {
            "meta": meta,
            "type": "logs",
            "entries": log_entries
        }
    elif kv_count >= len(non_empty_lines) * 0.5 and kv_count > 0:
        if len(sections) > 1 or (len(sections) == 1 and "general" not in sections):
            result = {
                "meta": meta,
                "type": "configuration",
                "sections": sections
            }
        else:
            result = {
                "meta": meta,
                "type": "key_value",
                "data": key_values
            }
    else:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        result = {
            "meta": meta,
            "type": "text_document",
            "paragraphs": paragraphs,
            "lines": lines
        }

    return json.dumps(result, ensure_ascii=False, indent=2)


def markdown_to_structured_json(md_text: str) -> str:
    """Parses Markdown structure into JSON with headings, sections, lists, and code blocks."""
    lines = md_text.splitlines()
    sections = []
    current_section = {
        "title": "Introduction",
        "level": 0,
        "paragraphs": [],
        "lists": [],
        "code_blocks": []
    }
    in_code = False
    code_lang = ""
    code_buffer = []
    current_list = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                current_section["code_blocks"].append({
                    "language": code_lang or "plain",
                    "code": "\n".join(code_buffer)
                })
                in_code = False
                code_buffer = []
            else:
                in_code = True
                code_lang = stripped[3:].strip()
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if stripped.startswith("#"):
            if current_list:
                current_section["lists"].append(current_list)
                current_list = []

            h_level = len(stripped.split()[0])
            h_text = stripped.lstrip("#").strip()
            if current_section["paragraphs"] or current_section["lists"] or current_section["code_blocks"] or current_section["level"] > 0:
                sections.append(current_section)
            current_section = {
                "title": h_text,
                "level": h_level,
                "paragraphs": [],
                "lists": [],
                "code_blocks": []
            }
            continue

        if stripped.startswith(("- ", "* ", "+ ")) or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in (". ", ") ")):
            item_text = re.sub(r'^(?:[-*+]|\d+[.)])\s+', '', stripped)
            current_list.append(item_text)
            continue
        elif current_list:
            current_section["lists"].append(current_list)
            current_list = []

        if stripped:
            current_section["paragraphs"].append(stripped)

    if current_list:
        current_section["lists"].append(current_list)
    sections.append(current_section)

    result = {
        "meta": {
            "total_sections": len(sections),
            "lines_count": len(lines),
            "words_count": len(md_text.split()),
        },
        "sections": sections
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def docx_to_structured_json(docx_bytes: bytes) -> str:
    """Extracts DOCX content into structured JSON containing headings, paragraphs, and tables."""
    doc = Document(io.BytesIO(docx_bytes))
    headings = []
    paragraphs = []
    tables_data = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        style = p.style.name.lower() if p.style else ""
        if "heading" in style or "title" in style:
            headings.append({"style": p.style.name if p.style else "Heading", "text": txt})
        paragraphs.append(txt)

    for t in doc.tables:
        if not t.rows:
            continue
        headers = [c.text.strip() for c in t.rows[0].cells]
        rows_list = []
        for r in t.rows[1:]:
            row_dict = {}
            for idx, c in enumerate(r.cells):
                h_name = headers[idx] if idx < len(headers) and headers[idx] else f"col_{idx+1}"
                row_dict[h_name] = _infer_value_type(c.text.strip())
            rows_list.append(row_dict)
        tables_data.append({
            "headers": headers,
            "rows": rows_list
        })

    result = {
        "meta": {
            "total_headings": len(headings),
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables_data),
        },
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables_data
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def csv_to_structured_json(csv_text: str) -> str:
    """Converts CSV/TSV to typed and structured JSON with table meta."""
    delimiter = "\t" if "\t" in csv_text and "," not in csv_text else (";" if ";" in csv_text and "," not in csv_text else ",")
    reader = csv.DictReader(io.StringIO(csv_text), delimiter=delimiter)
    typed_rows = []
    fieldnames = reader.fieldnames or []

    for row in reader:
        typed_row = {}
        for k, v in row.items():
            typed_row[k] = _infer_value_type(v) if v is not None else None
        typed_rows.append(typed_row)

    result = {
        "meta": {
            "columns": fieldnames,
            "total_rows": len(typed_rows)
        },
        "data": typed_rows
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


# ==========================================
# Readable HTML Converters
# ==========================================

def text_to_readable_html(text: str, title: str = "Документ") -> str:
    """Converts plain text or log into a beautifully formatted, readable HTML webpage."""
    lines = text.splitlines()
    non_empty = [l.strip() for l in lines if l.strip()]

    log_pattern = re.compile(r'^(\d{4}[-/.]\d{2}[-/.]\d{2}[ T]\d{2}:\d{2}:\d{2}(?:[.,]\d+)?)\s*(?:\[(\w+)\]|(\w+):)?\s*(.*)$')
    is_log = sum(1 for l in non_empty if log_pattern.match(l)) >= len(non_empty) * 0.6 if non_empty else False

    body_parts = []
    body_parts.append(f"<h1>{html.escape(title)}</h1>")

    if is_log:
        body_parts.append("<div class='log-container'>")
        for line in lines:
            m = log_pattern.match(line.strip())
            if m:
                ts, lvl1, lvl2, msg = m.groups()
                lvl = (lvl1 or lvl2 or "INFO").lower()
                badge_class = f"log-{lvl}" if lvl in ("info", "warn", "warning", "error", "debug") else "log-debug"
                body_parts.append(
                    f"<div class='log-row'>"
                    f"<span class='log-ts'>{html.escape(ts)}</span>"
                    f"<span class='log-badge {badge_class}'>{html.escape(lvl.upper())}</span>"
                    f"<span>{html.escape(msg)}</span>"
                    f"</div>"
                )
            elif line.strip():
                body_parts.append(f"<div class='log-row'>{html.escape(line)}</div>")
        body_parts.append("</div>")
    else:
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            para_clean = para.strip()
            if not para_clean:
                continue

            para_escaped = html.escape(para_clean)
            para_linked = re.sub(
                r'(https?://[^\s<>"]+|www\.[^\s<>"]+)',
                r'<a href="\1" target="_blank" rel="noopener">\1</a>',
                para_escaped
            )
            para_html = para_linked.replace("\n", "<br>")
            body_parts.append(f"<p>{para_html}</p>")

    return _wrap_in_readable_html_template("\n".join(body_parts), title=title)


def markdown_to_readable_html(md_text: str, title: str = "Документ") -> str:
    """Converts Markdown to a full-featured styled HTML document."""
    html_body = markdown.markdown(md_text, extensions=['extra', 'tables', 'fenced_code', 'nl2br', 'sane_lists'])
    return _wrap_in_readable_html_template(html_body, title=title)


def docx_to_readable_html(docx_bytes: bytes, title: str = "Документ") -> str:
    """Converts DOCX into a styled readable HTML document."""
    md_text = docx_to_markdown(docx_bytes)
    return markdown_to_readable_html(md_text, title=title)


def csv_to_readable_html(csv_text: str, title: str = "Таблица") -> str:
    """Converts CSV table into an interactive, readable HTML table."""
    md_table = csv_to_markdown(csv_text)
    return markdown_to_readable_html(md_table, title=title)


# ==========================================
# DOCX <-> Markdown Core Handlers
# ==========================================

def _add_markdown_runs_to_paragraph(paragraph, text: str) -> None:
    pattern = re.compile(r'(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|___[^_]+___|__[^_]+__|_[^_]+_)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
        elif (part.startswith("***") and part.endswith("***") and len(part) >= 6) or \
             (part.startswith("___") and part.endswith("___") and len(part) >= 6):
            r = paragraph.add_run(part[3:-3])
            r.bold = True
            r.italic = True
        elif (part.startswith("**") and part.endswith("**") and len(part) >= 4) or \
             (part.startswith("__") and part.endswith("__") and len(part) >= 4):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif (part.startswith("*") and part.endswith("*") and len(part) >= 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) >= 2):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text: str) -> bytes:
    """Converts Markdown text into a styled DOCX document."""
    doc = Document()
    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph("\n".join(code_lines))
                for run in p.runs:
                    run.font.name = "Courier New"
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith(("- ", "* ", "+ ")):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in (". ", ") "):
            doc.add_paragraph(stripped[3:], style='List Number')
        elif stripped.startswith("> "):
            doc.add_paragraph(stripped[2:], style='Quote' if 'Quote' in doc.styles else 'Normal')
        elif stripped == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            _add_markdown_runs_to_paragraph(p, line)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def docx_to_markdown(docx_bytes: bytes) -> str:
    """Converts a DOCX document into Markdown string."""
    doc = Document(io.BytesIO(docx_bytes))
    md_parts = []

    for p in doc.paragraphs:
        style_name = p.style.name.lower() if p.style else ""
        text = ""
        for r in p.runs:
            run_text = r.text
            if not run_text:
                continue
            if r.bold and r.italic:
                run_text = f"***{run_text}***"
            elif r.bold:
                run_text = f"**{run_text}**"
            elif r.italic:
                run_text = f"*{run_text}*"
            text += run_text

        if not text.strip():
            md_parts.append("")
            continue

        if "heading 1" in style_name:
            md_parts.append(f"# {text.strip()}")
        elif "heading 2" in style_name:
            md_parts.append(f"## {text.strip()}")
        elif "heading 3" in style_name:
            md_parts.append(f"### {text.strip()}")
        elif "heading 4" in style_name:
            md_parts.append(f"#### {text.strip()}")
        elif "bullet" in style_name or "list bullet" in style_name:
            md_parts.append(f"- {text.strip()}")
        elif "number" in style_name or "list number" in style_name:
            md_parts.append(f"1. {text.strip()}")
        elif "quote" in style_name:
            md_parts.append(f"> {text.strip()}")
        else:
            md_parts.append(text)

    for table in doc.tables:
        if not table.rows:
            continue
        table_lines = []
        headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        table_lines.append("| " + " | ".join(headers) + " |")
        table_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
        md_parts.append("\n" + "\n".join(table_lines) + "\n")

    return "\n".join(md_parts)


def docx_to_text(docx_bytes: bytes) -> str:
    """Extracts plain text from DOCX."""
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def text_to_docx(text: str) -> bytes:
    """Converts plain text to DOCX."""
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def html_to_markdown(html_text: str) -> str:
    """HTML to Markdown converter."""
    text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n\n', html_text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<strong[^>]*>(.*?)</strong>|<b[^>]*>(.*?)</b>', r'**\1\2**', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<em[^>]*>(.*?)</em>|<i[^>]*>(.*?)</i>', r'*\1\2*', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def html_to_text(html_text: str) -> str:
    """Strips HTML tags and extracts text."""
    clean = re.sub(r'<br\s*/?>', '\n', html_text, flags=re.IGNORECASE)
    clean = re.sub(r'</p>', '\n\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'</h1>|</h2>|</h3>|</h4>', '\n\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'</li>', '\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'<[^>]+>', '', clean)
    return clean.strip()


def json_to_csv(text: str) -> str:
    """Converts JSON to CSV string."""
    data = json.loads(text)
    out = io.StringIO()
    if isinstance(data, dict) and "data" in data and isinstance(data["data"], list):
        data = data["data"]

    if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
        keys = list(data[0].keys())
        writer = csv.DictWriter(out, fieldnames=keys)
        writer.writeheader()
        writer.writerows(data)
    elif isinstance(data, dict):
        writer = csv.writer(out)
        writer.writerow(["Key", "Value"])
        for k, v in data.items():
            val = json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else v
            writer.writerow([k, val])
    elif isinstance(data, list):
        writer = csv.writer(out)
        writer.writerow(["Item"])
        for item in data:
            writer.writerow([item])
    else:
        writer = csv.writer(out)
        writer.writerow(["Value"])
        writer.writerow([data])
    return out.getvalue()


def csv_to_markdown(text: str) -> str:
    """Converts CSV to Markdown table."""
    delimiter = "\t" if "\t" in text and "," not in text else (";" if ";" in text and "," not in text else ",")
    reader = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    if not reader:
        return text
    md_lines = []
    headers = [cell.strip().replace("\n", " ") for cell in reader[0]]
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in reader[1:]:
        row_cells = [c.strip().replace("\n", " ") for c in row]
        if len(row_cells) < len(headers):
            row_cells.extend([""] * (len(headers) - len(row_cells)))
        md_lines.append("| " + " | ".join(row_cells[:len(headers)]) + " |")
    return "\n".join(md_lines)


def csv_to_docx(text: str) -> bytes:
    """Converts CSV table to Word DOCX document."""
    delimiter = "\t" if "\t" in text and "," not in text else (";" if ";" in text and "," not in text else ",")
    reader = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    doc = Document()
    if not reader:
        doc.add_paragraph("Пустая таблица")
    else:
        cols_count = max(len(r) for r in reader) if reader else 1
        table = doc.add_table(rows=len(reader), cols=cols_count)
        table.style = 'Table Grid'
        for i, row in enumerate(reader):
            for j, cell_text in enumerate(row):
                if j < cols_count:
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    if i == 0:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


# ==========================================
# Main Document Converter Router
# ==========================================

def convert_document(input_bytes: bytes, source_format: str, target_format: str) -> tuple[bytes, str]:
    """
    Converts document or text between formats.
    Returns (output_bytes, file_extension).
    """
    src = normalize_format(source_format)
    target = normalize_format(target_format)

    # 1. DOCX source
    if src == "DOCX":
        if target == "MD":
            md_str = docx_to_markdown(input_bytes)
            return md_str.encode("utf-8"), "md"
        elif target == "JSON":
            json_str = docx_to_structured_json(input_bytes)
            return json_str.encode("utf-8"), "json"
        elif target == "HTML":
            html_str = docx_to_readable_html(input_bytes)
            return html_str.encode("utf-8"), "html"
        elif target in ("TXT", "DAT", "LOG"):
            txt_str = docx_to_text(input_bytes)
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return txt_str.encode("utf-8"), ext
        else:
            txt_str = docx_to_text(input_bytes)
            ext = SUPPORTED_DOCUMENT_FORMATS.get(target, {}).get("ext", target.lower())
            return txt_str.encode("utf-8"), ext

    # Non-DOCX source: decode text
    text_content = decode_text(input_bytes)

    # 2. Target JSON
    if target == "JSON":
        if src == "MD":
            json_str = markdown_to_structured_json(text_content)
        elif src in ("CSV", "TSV"):
            json_str = csv_to_structured_json(text_content)
        else:
            json_str = text_to_structured_json(text_content)
        return json_str.encode("utf-8"), "json"

    # 3. Target HTML
    if target == "HTML":
        if src == "MD":
            html_str = markdown_to_readable_html(text_content)
        elif src in ("CSV", "TSV"):
            html_str = csv_to_readable_html(text_content)
        elif src == "JSON":
            try:
                pretty = json.dumps(json.loads(text_content), ensure_ascii=False, indent=2)
            except Exception:
                pretty = text_content
            body = f"<h1>JSON Документ</h1><pre><code>{html.escape(pretty)}</code></pre>"
            html_str = _wrap_in_readable_html_template(body, title="JSON Документ")
        else:
            html_str = text_to_readable_html(text_content)
        return html_str.encode("utf-8"), "html"

    # 4. Target DOCX
    if target == "DOCX":
        if src == "MD":
            return markdown_to_docx(text_content), "docx"
        elif src in ("CSV", "TSV"):
            return csv_to_docx(text_content), "docx"
        elif src == "HTML":
            md_from_html = html_to_markdown(text_content)
            return markdown_to_docx(md_from_html), "docx"
        else:
            return text_to_docx(text_content), "docx"

    # 5. MD as source
    if src == "MD":
        if target in ("TXT", "DAT", "LOG"):
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return text_content.encode("utf-8"), ext

    # 6. Target MD
    if target == "MD":
        if src in ("CSV", "TSV"):
            return csv_to_markdown(text_content).encode("utf-8"), "md"
        elif src == "HTML":
            return html_to_markdown(text_content).encode("utf-8"), "md"
        elif src == "JSON":
            try:
                parsed = json.loads(text_content)
                pretty_json = json.dumps(parsed, ensure_ascii=False, indent=2)
                return f"```json\n{pretty_json}\n```\n".encode("utf-8"), "md"
            except Exception:
                return text_content.encode("utf-8"), "md"
        else:
            return text_content.encode("utf-8"), "md"

    # 7. CSV / TSV
    if src == "CSV":
        if target == "TSV":
            tsv_str = text_content.replace(",", "\t")
            return tsv_str.encode("utf-8"), "tsv"
        elif target in ("TXT", "DAT", "LOG"):
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return text_content.encode("utf-8"), ext

    if src == "TSV":
        if target == "CSV":
            csv_str = text_content.replace("\t", ",")
            return csv_str.encode("utf-8"), "csv"
        elif target in ("TXT", "DAT", "LOG"):
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return text_content.encode("utf-8"), ext

    # 8. JSON
    if src == "JSON":
        if target == "CSV":
            return json_to_csv(text_content).encode("utf-8"), "csv"
        elif target in ("TXT", "DAT", "LOG"):
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return text_content.encode("utf-8"), ext

    # 9. HTML
    if src == "HTML":
        if target in ("TXT", "DAT", "LOG"):
            plain = html_to_text(text_content)
            ext = SUPPORTED_DOCUMENT_FORMATS[target]["ext"]
            return plain.encode("utf-8"), ext

    # 10. General fallback for text files
    ext = SUPPORTED_DOCUMENT_FORMATS.get(target, {}).get("ext", target.lower())
    return text_content.encode("utf-8"), ext
